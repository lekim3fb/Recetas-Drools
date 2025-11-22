from docx import Document
from docx.shared import Pt
import os

# Build the report content
project_root = os.path.abspath(os.path.dirname(__file__))
output_path = os.path.join(project_root, 'Plantilla_Report.docx')

doc = Document()

# Helper styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
doc.add_heading('Informe del Sistema de Recomendacion de Recetas', level=1)

# Description of the system
doc.add_heading('Descripcion del sistema', level=2)
doc.add_paragraph(
    'Sistema de recomendacion de recetas basado en Drools (motor de reglas). ' 
    'Proporciona recomendaciones de recetas personalizadas a partir de: datos del usuario (ingredientes disponibles, preferencias, restricciones, tiempo y nivel de habilidad), ' 
    'una base de datos inicial de ingredientes y recetas, reglas moduladas en archivos DRL, y una pequena integracion DMN para evaluacion de la calidad de la receta.'
)

# Conceptualization
doc.add_heading('La conceptualizacion', level=2)
doc.add_paragraph(
    'El sistema se diseno siguiendo un enfoque de Sistemas Basados en Conocimiento (SBC). ' 
    'Se definieron las siguientes capas: declarativa (entidades Java: Ingrediente, Receta, IngredienteReceta, Usuario, Recomendacion, EvaluacionReceta), ' 
    'procedimental (reglas Drools divididas en modulos: inicializacion de datos, validacion, filtrado, puntuacion, explicacion e informe), y heuristica (bonificaciones y penalizaciones para el ranking).' 
)

doc.add_paragraph('Fuentes y evidencias utilizadas para la conceptualizacion:')
para = doc.add_paragraph()
para.style = 'List Bullet'
para.add_run('Documentacion del proyecto y requisitos provistos por el cliente/usuario.')
para = doc.add_paragraph()
para.style = 'List Bullet'
para.add_run('Buenas practicas de diseno SBC y patrones de integracion Drools + DMN.')
para = doc.add_paragraph()
para.style = 'List Bullet'
para.add_run('Contenido y reglas desarrolladas iterativamente en archivos DRL del proyecto.')

# Formalization reasoned including control part
doc.add_heading('Formalizacion razonada (incluyendo control)', level=2)
doc.add_paragraph(
    'La formalizacion convierte la conceptualizacion a artefactos ejecutables: clases Java que representan el dominio y modulos DRL que implementan el ciclo de decision. ' 
    'Se aplico modularizacion en la base de conocimiento (kmodules) para separar responsabilidades: 01_basedatos (inicializa hechos), 02_validacion (verifica datos del usuario), 03_filtrado (descarta recetas por alergia/tiempo/dieta/dificultad), 04_puntuacion (calcula puntuaciones y aplica bonificaciones), 05_explicacion (genera justificaciones), 06_informe (presenta resultados) y 07_decision_dmn (integra DMN para evaluacion adicional).' 
)

doc.add_paragraph(
    'Parte de control: el flujo es impulsado por la insercion del Usuario en la sesion Kie. Una vez disponibles los hechos (ingredientes y recetas inicializadas desde 01_basedatos), las reglas de validacion se aplican, luego las de filtrado, puntuacion y finalmente la generacion de explicaciones e informe. El control es implicito en el motor de reglas (orden por salience y activacion de reglas), y explicito en el codigo Java que decide cuando crear la sesion, insertar el usuario y disparar fireAllRules().'
)

doc.add_heading('Presentacion de resultados', level=2)
doc.add_paragraph(
    'El sistema de presentacion de resultados ha sido diseado para mostrar unicamente informacion relevante al usuario final. '
    'La interfaz interactiva en Java (Main.java) recopila los datos del usuario mediante un flujo guiado con validaciones robustas '
    '(re-solicita entrada hasta que sea valida para opciones como nivel de habilidad, ingredientes, preferencias y restricciones). '
    'Una vez procesadas las reglas, la salida muestra:'
)

p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('Recetas recomendadas: nombre, estado, puntuacion y razones de recomendacion.')

p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('Ranking ordenado por puntuacion de las recetas viables.')

p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('Si no existen recomendaciones: mensaje claro indicando que con las caracteristicas del usuario no hay una receta buena para el.')

doc.add_paragraph(
    'Las recetas descartadas no se muestran en la salida final para evitar informacion negativa innecesaria al usuario. '
    'El enfoque es presentar opciones constructivas: o mostramos lo que si le recomendamos, o le informamos de forma amable que no hay receta adecuada en la base de datos actual.'
)

# Conclusions: difficulties, time spent, comments
doc.add_heading('Conclusiones', level=2)
doc.add_paragraph('Dificultades encontradas:')

p = doc.add_paragraph()
p.style = 'List Number'
p.add_run('Compatibilidades de sintaxis con Drools 8:')
p = doc.add_paragraph('Se corrigieron usos invalidos como order by en queries y llamadas inexistentes como drools.getObjectsInMemory().', style='List Bullet')

p = doc.add_paragraph()
p = doc.add_paragraph()
p.style = 'List Number'
p.add_run('Eval y expresiones complejas:')
p = doc.add_paragraph('Sustituir eval multi-linea por expresiones compatibles y patrones de hechos para evitar errores de parsing.', style='List Bullet')

p = doc.add_paragraph()
p = doc.add_paragraph()
p.style = 'List Number'
p.add_run('Integracion y pruebas en terminal:')
p = doc.add_paragraph('Asegurar que los comandos Maven se ejecuten en el directorio raiz del proyecto; se detectaron problemas al ejecutar mvn desde subdirectorios.', style='List Bullet')

# Time spent estimation
doc.add_heading('Tiempo empleado (estimacion por apartados)', level=2)

doc.add_paragraph('Resumen estimado: ~22-24 horas (depende de pruebas y revisiones adicionales).')

# Comments
doc.add_heading('Comentarios finales', level=2)
doc.add_paragraph(
    'El proyecto proporciona una base robusta para un sistema de recomendacion declarativo usando Drools. '
    'Se ha optimizado la experiencia del usuario ocultando informacion negativa (recetas descartadas) y mostrando un mensaje amable cuando no se encuentra una recomendacion adecuada. '
    'Se recomienda: (1) anadir tests unitarios y de integracion para las reglas, (2) extraer la inicializacion de datos a ficheros de recursos o una DB para facilidad de mantenimiento, '
    '(3) mejorar el manejo de internacionalizacion/encodings para prevenir caracteres extraños en la consola, y (4) anadir logging (SLF4J) con implementacion para produccion.'
)

# Save document
try:
    doc.save(output_path)
    print('OK: Archivo generado ->', output_path)
except Exception as e:
    print('ERROR al generar el documento:', e)
