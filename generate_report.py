from docx import Document
from docx.shared import Pt
import os
import shutil
import time

# Build the report content based on the project and conversation summary
project_root = os.path.abspath(os.path.dirname(__file__))
output_path = os.path.join(project_root, 'Plantilla_Report.docx')
temp_path = os.path.join(project_root, 'Plantilla_Report_TEMP.docx')

doc = Document()

# Helper styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
doc.add_heading('Informe del Sistema de Recomendación de Recetas', level=1)

# Description of the system
doc.add_heading('Descripción del sistema', level=2)
doc.add_paragraph(
    'Sistema de recomendación de recetas basado en Drools (motor de reglas). ' 
    'Proporciona recomendaciones de recetas personalizadas a partir de: datos del usuario (ingredientes disponibles, preferencias, restricciones, tiempo y nivel de habilidad), ' 
    'una base de datos inicial de ingredientes y recetas, reglas moduladas en archivos DRL, y una pequeña integración DMN para evaluación de la calidad de la receta.'
)

# Conceptualization
doc.add_heading('La conceptualización', level=2)
doc.add_paragraph(
    'El sistema se diseñó siguiendo un enfoque de Sistemas Basados en Conocimiento (SBC). ' 
    'Se definieron las siguientes capas: declarativa (entidades Java: Ingrediente, Receta, IngredienteReceta, Usuario, Recomendacion, EvaluacionReceta), ' 
    'procedimental (reglas Drools divididas en módulos: inicialización de datos, validación, filtrado, puntuación, explicación e informe), y heurística (bonificaciones y penalizaciones para el ranking).' 
)

doc.add_paragraph('Fuentes y evidencias utilizadas para la conceptualización:')
para = doc.add_paragraph()
para.style = 'List Bullet'
para.add_run('Documentación del proyecto y requisitos provistos por el cliente/usuario.').bold = False
para = doc.add_paragraph()
para.style = 'List Bullet'
para.add_run('Buenas prácticas de diseño SBC y patrones de integración Drools + DMN.').bold = False
para = doc.add_paragraph()
para.style = 'List Bullet'
para.add_run('Contenido y reglas desarrolladas iterativamente en archivos DRL del proyecto.').bold = False

# Formalization reasoned including control part
doc.add_heading('Formalización razonada (incluyendo control)', level=2)
doc.add_paragraph(
    'La formalización convierte la conceptualización a artefactos ejecutables: clases Java que representan el dominio y módulos DRL que implementan el ciclo de decisión. ' 
    'Se aplicó modularización en la base de conocimiento (kmodules) para separar responsabilidades: 01_basedatos (inicializa hechos), 02_validacion (verifica datos del usuario), 03_filtrado (descarta recetas por alergia/tiempo/dieta/dificultad), 04_puntuacion (calcula puntuaciones y aplica bonificaciones), 05_explicacion (genera justificaciones), 06_informe (presenta resultados) y 07_decision_dmn (integra DMN para evaluación adicional).' 
)

doc.add_paragraph(
    """Parte de control: el flujo es impulsado por la inserción del Usuario en la sesión Kie. Una vez disponibles los hechos (ingredientes y recetas inicializadas desde 01_basedatos), las reglas de validación se aplican, luego las de filtrado, puntuación y finalmente la generación de explicaciones e informe. El control es implícito en el motor de reglas (orden por salience y activación de reglas), y explícito en el código Java que decide cuándo crear la sesión, insertar el usuario y disparar fireAllRules()."""
)

doc.add_heading('Presentación de resultados', level=2)
doc.add_paragraph(
    'El sistema de presentación de resultados ha sido diseñado para mostrar únicamente información relevante al usuario final. '
    'La interfaz interactiva en Java (Main.java) recopila los datos del usuario mediante un flujo guiado con validaciones robustas '
    '(re-solicita entrada hasta que sea válida para opciones como nivel de habilidad, ingredientes, preferencias y restricciones). '
    'Una vez procesadas las reglas, la salida muestra:'
)

p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('Recetas recomendadas: nombre, estado, puntuación y razones de recomendación.')

p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('Ranking ordenado por puntuación de las recetas viables.')

p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('Si no existen recomendaciones: mensaje claro "Con las características del usuario no hay una receta buena para él."')

doc.add_paragraph(
    'Las recetas descartadas no se muestran en la salida final para evitar información negativa innecesaria al usuario. '
    'El enfoque es presentar opciones constructivas: o mostramos lo que sí le recomendamos, o le informamos de forma amable que no hay receta adecuada en la base de datos actual.'
)

# Conclusions: difficulties, time spent, comments
doc.add_heading('Conclusiones', level=2)
doc.add_paragraph('Dificultades encontradas:')

p = doc.add_paragraph()
p.style = 'List Number'
p.add_run('Compatibilidades de sintaxis con Drools 8:').bold = True
p = doc.add_paragraph('   - Se corrigieron usos inválidos como `order by` en queries y llamadas inexistentes como `drools.getObjectsInMemory()`.', style='List Bullet')

p = doc.add_paragraph()
p = doc.add_paragraph()
p.style = 'List Number'
p.add_run('Eval y expresiones complejas:').bold = True
p = doc.add_paragraph('   - Sustituir eval multi-línea por expresiones compatibles y patrones de hechos para evitar errores de parsing.', style='List Bullet')

p = doc.add_paragraph()
p = doc.add_paragraph()
p.style = 'List Number'
p.add_run('Integración y pruebas en terminal:').bold = True
p = doc.add_paragraph('   - Asegurar que los comandos Maven se ejecuten en el directorio raíz del proyecto; se detectaron problemas al ejecutar mvn desde subdirectorios.', style='List Bullet')

# Time spent estimation
doc.add_heading('Tiempo empleado (estimación por apartados)', level=2)
rows = [
    ('Análisis y diseño conceptual', '~4 horas'),
    ('Modelado de entidades Java', '~2 horas'),
    ('Desarrollo y modularización de reglas DRL', '~6-8 horas'),
    ('Correcciones sintácticas y compatibilidad con Drools 8', '~3 horas'),
    ('Integración DMN y pruebas', '~2 horas'),
    ('Implementación de interfaz interactiva (Main)', '~2 horas'),
    ('Pruebas de ejecución y depuración', '~3 horas')
]

doc.add_paragraph('Resumen estimado: ~22-24 horas (depende de pruebas y revisiones adicionales).')

# Comments
doc.add_heading('Comentarios finales', level=2)
doc.add_paragraph(
    'El proyecto proporciona una base robusta para un sistema de recomendación declarativo usando Drools. '
    'Se ha optimizado la experiencia del usuario ocultando información negativa (recetas descartadas) y mostrando un mensaje amable cuando no se encuentra una recomendación adecuada. '
    'Se recomienda: (1) añadir tests unitarios y de integración para las reglas, (2) extraer la inicialización de datos a ficheros de recursos o una DB para facilidad de mantenimiento, '
    '(3) mejorar el manejo de internacionalización/encodings para prevenir caracteres extraños en la consola, y (4) añadir logging (SLF4J) con implementación para producción.'
)

# Save document
try:
    doc.save(temp_path)
    # Try to remove old file
    max_retries = 3
    for attempt in range(max_retries):
        try:
            if os.path.exists(output_path):
                os.remove(output_path)
            shutil.move(temp_path, output_path)
            print('OK: Archivo generado ->', output_path)
            break
        except PermissionError:
            if attempt < max_retries - 1:
                time.sleep(1)
            else:
                # If all retries failed, just save to temp and inform user
                print('OK: Archivo generado (temp) ->', temp_path)
                print('Note: El archivo original esta en uso. Temp file:', temp_path)
except Exception as e:
    print('ERROR al generar el documento:', e)
